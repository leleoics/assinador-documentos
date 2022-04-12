import pandas as pd
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import win32com.client
from PyPDF2 import PdfFileMerger, PdfFileReader

def clear_prompt(): # Função para limpar prompt de comando
    os.system('cls' if os.name == 'nt' else 'clear')


print("Selecione a opção desejada:")
print("1- Preencher e gerar arquivos PDF")
print("2- Sair")
mergedObject = PdfFileMerger()
while True:
    try:
        choice = int(input("Digite a opção desejada: "))
        if not 1 <= choice <= 2:
            raise ValueError("Opção não encontrada")
    except ValueError as e:
        print("Digite uma das opções!")
        print("Exemplo: 1")
    if choice == 1:
        df = pd.read_excel("./1- Arquivos de entrada/Lista.xlsx")
        tamanho = len(df)

        for i in range(0,len(df)):
            f = open("./1- Arquivos de entrada/Documento.docx", 'rb')
            document = Document(f)
            nome = str(df['Nome'][i])
            matricula = str(df['Matricula'][i])
            contador = i + 1
            resto = tamanho - contador
            print(f'Preenchendo o {contador}º arquivo, de {tamanho}.')
            print(f'Restando {resto} arquivos...')
            
            for paragraph in document.paragraphs:
                paragraph.add_run()
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.line_spacing = 1.75
                if '{nome}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{nome}',nome)
                if '{matricula}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{matricula}', matricula)
            
            nome_arquivo = "Página - " + str(contador) + ".pdf"
            os.chdir("./3- Dados auxiliares")
            document.save('Arquivo.docx')
            f.close()
            wdFormatPDF = 17
            outputFile = os.path.abspath(nome_arquivo)
            inputFile = os.path.abspath('Arquivo.docx')
            os.chdir("..")
            word = win32com.client.Dispatch('Word.Application')
            doc = word.Documents.Open(inputFile)
            doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()
            del document, doc
        print("Documentos Preenchidos!")
        print("Mesclando documentos...")
        os.chdir("./3- Dados auxiliares")
        for fileNumber in range(1, (len(df)+1)):
            if fileNumber == (len(df)):
                mergedObject.append(PdfFileReader("Página - " + str(fileNumber)+ ".pdf", 'rb'))
                os.chdir("..")
                os.chdir("./2- Arquivos gerados")
                mergedObject.write("Desenvolvimento Pessoal.pdf")
                break
            else:
                resto = len(df) - fileNumber
                print(f"Restando {resto} documentos para mesclar.")
                mergedObject.append(PdfFileReader("Página - " + str(fileNumber)+ ".pdf", 'rb'))


    if choice == 2:
        print("Operação cancelada")
        break
    else:
         break

input('Precione enter para encerrar o programa')

